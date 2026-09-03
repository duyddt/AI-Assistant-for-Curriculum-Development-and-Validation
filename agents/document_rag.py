"""Document ingestion and grounded retrieval for the SV5 prototype.

The module deliberately keeps document parsing and retrieval separate from the
LLM.  This makes the source evidence testable without an API key and lets the
production service replace :class:`DocumentRetriever` with the RAG service of
SV1 (for example Qdrant + bge-m3) later without changing the SV5 contract.

Supported prototype formats: PDF, DOCX, Markdown and plain text.  PDF chunks
retain page numbers; DOCX chunks retain paragraph/table locations because a
DOCX file has no reliable physical page number before rendering.
"""

from __future__ import annotations

import hashlib
import math
import re
from collections import Counter
from pathlib import Path
from typing import Iterable, Sequence

from pydantic import BaseModel, Field


class DocumentIngestionError(ValueError):
    """Raised when a source document is missing, unsupported, or unreadable."""


class SourceChunk(BaseModel):
    """An auditable text chunk extracted from exactly one source location."""

    document_id: str
    source_file: str
    chunk_id: str
    location: str
    text: str
    page: int | None = None


class SearchHit(BaseModel):
    chunk: SourceChunk
    score: float


class SourceCitation(BaseModel):
    """Evidence returned to callers.  ``used_for`` shows which output it backs."""

    source_file: str
    chunk_id: str
    location: str
    excerpt: str
    page: int | None = None
    score: float = 0.0
    used_for: list[str] = Field(default_factory=list)


class DocumentIngestionSummary(BaseModel):
    source_files: list[str]
    chunk_count: int
    chunks_by_source: dict[str, int] = Field(default_factory=dict)
    retrieval_strategy: str = "lexical_tfidf"


_SUPPORTED_SUFFIXES = {".pdf", ".docx", ".md", ".txt"}
_TOKEN_PATTERN = re.compile(r"[\wÀ-ỹ]+", re.UNICODE)


def _tokenize(text: str) -> list[str]:
    return [token.lower() for token in _TOKEN_PATTERN.findall(text) if len(token) >= 2]


def _normalise_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    """Split a unit of text while retaining a small overlap for context."""
    text = _normalise_whitespace(text)
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    parts: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        if end < len(text):
            split_at = text.rfind(" ", start + chunk_size // 2, end)
            if split_at > start:
                end = split_at
        part = text[start:end].strip()
        if part:
            parts.append(part)
        if end >= len(text):
            break
        start = max(end - overlap, start + 1)
    return parts


def _pdf_units(path: Path) -> Iterable[tuple[str, str, int | None]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency checked by install
        raise DocumentIngestionError("Thiếu thư viện pypdf để đọc PDF") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001 - normalise parser errors at boundary
        raise DocumentIngestionError(f"Không đọc được PDF '{path.name}': {exc}") from exc

    found_text = False
    for page_number, page in enumerate(reader.pages, start=1):
        text = _normalise_whitespace(page.extract_text() or "")
        if text:
            found_text = True
            yield text, f"page {page_number}", page_number
    if not found_text:
        raise DocumentIngestionError(
            f"PDF '{path.name}' không có văn bản trích xuất được; cần OCR trước khi ingest"
        )


def _docx_units(path: Path) -> Iterable[tuple[str, str, int | None]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency checked by install
        raise DocumentIngestionError("Thiếu thư viện python-docx để đọc DOCX") from exc

    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001 - normalise parser errors at boundary
        raise DocumentIngestionError(f"Không đọc được DOCX '{path.name}': {exc}") from exc

    found_text = False
    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        text = _normalise_whitespace(paragraph.text)
        if text:
            found_text = True
            yield text, f"paragraph {paragraph_number}", None

    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            values = [_normalise_whitespace(cell.text) for cell in row.cells]
            text = " | ".join(value for value in values if value)
            if text:
                found_text = True
                yield text, f"table {table_number}, row {row_number}", None

    if not found_text:
        raise DocumentIngestionError(f"DOCX '{path.name}' không có văn bản trích xuất được")


def _text_units(path: Path) -> Iterable[tuple[str, str, int | None]]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8-sig")
    except OSError as exc:
        raise DocumentIngestionError(f"Không đọc được file '{path.name}': {exc}") from exc

    if not _normalise_whitespace(text):
        raise DocumentIngestionError(f"File '{path.name}' rỗng")
    yield text, "text", None


def extract_document_chunks(
    document_paths: Sequence[str | Path],
    *,
    chunk_size: int = 900,
    overlap: int = 140,
) -> list[SourceChunk]:
    """Extract auditable chunks from a list of uploaded/source documents."""
    if not document_paths:
        raise DocumentIngestionError("Cần ít nhất một tài liệu nguồn")
    if chunk_size <= overlap or overlap < 0:
        raise ValueError("chunk_size phải lớn hơn overlap và overlap không âm")

    chunks: list[SourceChunk] = []
    for raw_path in document_paths:
        path = Path(raw_path)
        if not path.is_file():
            raise DocumentIngestionError(f"Không tìm thấy tài liệu: {path}")
        suffix = path.suffix.lower()
        if suffix not in _SUPPORTED_SUFFIXES:
            raise DocumentIngestionError(
                f"Định dạng '{suffix or '(không có đuôi)'}' chưa được hỗ trợ: {path.name}"
            )

        if suffix == ".pdf":
            units = _pdf_units(path)
        elif suffix == ".docx":
            units = _docx_units(path)
        else:
            units = _text_units(path)

        document_id = hashlib.sha256(str(path.resolve()).encode("utf-8")).hexdigest()[:12]
        part_number = 0
        for text, location, page in units:
            for part in _split_text(text, chunk_size, overlap):
                part_number += 1
                chunks.append(
                    SourceChunk(
                        document_id=document_id,
                        source_file=path.name,
                        chunk_id=f"{document_id}:{part_number:04d}",
                        location=location,
                        page=page,
                        text=part,
                    )
                )

    if not chunks:
        raise DocumentIngestionError("Không trích xuất được đoạn văn bản nào từ tài liệu nguồn")
    return chunks


class DocumentRetriever:
    """A transparent in-memory lexical retriever for the SV5 prototype.

    It is intentionally deterministic so unit tests can verify the evidence.
    The production integration point is this class: it may later delegate to
    SV1's hybrid/vector RAG API without changing downstream SV5 code.
    """

    retrieval_strategy = "lexical_tfidf"

    def __init__(self, chunks: Sequence[SourceChunk]):
        if not chunks:
            raise DocumentIngestionError("Không thể tạo retriever từ danh sách chunk rỗng")
        self.chunks = list(chunks)
        self._chunk_terms = [Counter(_tokenize(chunk.text)) for chunk in self.chunks]
        self._doc_frequency: Counter[str] = Counter()
        for terms in self._chunk_terms:
            self._doc_frequency.update(terms.keys())
        self._avg_length = sum(sum(terms.values()) for terms in self._chunk_terms) / len(self.chunks)

    @classmethod
    def from_paths(cls, document_paths: Sequence[str | Path]) -> "DocumentRetriever":
        return cls(extract_document_chunks(document_paths))

    def summary(self) -> DocumentIngestionSummary:
        chunks_by_source = Counter(chunk.source_file for chunk in self.chunks)
        return DocumentIngestionSummary(
            source_files=sorted({chunk.source_file for chunk in self.chunks}),
            chunk_count=len(self.chunks),
            chunks_by_source=dict(sorted(chunks_by_source.items())),
            retrieval_strategy=self.retrieval_strategy,
        )

    def search(self, query: str, limit: int = 3) -> list[SearchHit]:
        if limit < 1:
            raise ValueError("limit phải lớn hơn 0")
        query_terms = _tokenize(query)
        if not query_terms:
            return []

        unique_terms = set(query_terms)
        total_chunks = len(self.chunks)
        normalised_query = " ".join(query_terms)
        hits: list[SearchHit] = []
        for chunk, terms in zip(self.chunks, self._chunk_terms):
            length = max(1, sum(terms.values()))
            score = 0.0
            for term in unique_terms:
                frequency = terms.get(term, 0)
                if not frequency:
                    continue
                inverse_document_frequency = math.log((total_chunks + 1) / (self._doc_frequency[term] + 1)) + 1
                # BM25-style term saturation, without an external dependency.
                score += inverse_document_frequency * (
                    frequency * 2.2 / (frequency + 1.2 * (0.25 + 0.75 * length / self._avg_length))
                )

            normalised_text = " ".join(_tokenize(chunk.text))
            if normalised_query and normalised_query in normalised_text:
                score += 3.0
            if score > 0:
                hits.append(SearchHit(chunk=chunk, score=round(score, 5)))

        return sorted(hits, key=lambda hit: hit.score, reverse=True)[:limit]


def retrieve_sv5_evidence(
    document_paths: Sequence[str | Path],
    *,
    course_name: str,
    clo_descriptions: Sequence[str],
    topics: Sequence[str],
    limit_per_query: int = 3,
) -> tuple[DocumentIngestionSummary, list[SourceCitation]]:
    """Retrieve traceable source evidence for the three LLM tasks in SV5."""
    retriever = DocumentRetriever.from_paths(document_paths)
    queries = {
        "course_summary": " ".join([course_name, *clo_descriptions]),
        "weekly_schedule": " ".join([course_name, *topics, *clo_descriptions]),
        "rubrics": " ".join([course_name, *clo_descriptions]),
    }

    citations_by_chunk: dict[str, SourceCitation] = {}
    for used_for, query in queries.items():
        for hit in retriever.search(query, limit=limit_per_query):
            existing = citations_by_chunk.get(hit.chunk.chunk_id)
            if existing is None:
                existing = SourceCitation(
                    source_file=hit.chunk.source_file,
                    chunk_id=hit.chunk.chunk_id,
                    location=hit.chunk.location,
                    page=hit.chunk.page,
                    excerpt=hit.chunk.text[:800],
                    score=hit.score,
                    used_for=[used_for],
                )
                citations_by_chunk[hit.chunk.chunk_id] = existing
            elif used_for not in existing.used_for:
                existing.used_for.append(used_for)
                existing.score = max(existing.score, hit.score)

    citations = sorted(citations_by_chunk.values(), key=lambda item: item.score, reverse=True)
    if not citations:
        raise DocumentIngestionError(
            "Không tìm thấy đoạn tài liệu liên quan tới học phần/CLO/chủ đề; không sinh kết quả không có căn cứ"
        )
    return retriever.summary(), citations


def render_evidence_context(citations: Sequence[SourceCitation], max_characters: int = 6000) -> str:
    """Render untrusted source text safely for an LLM prompt."""
    if not citations:
        return "Không có tài liệu nguồn được cung cấp."

    entries: list[str] = []
    used = 0
    for citation in citations:
        locator = f"trang {citation.page}" if citation.page is not None else citation.location
        header = f"[Nguồn: {citation.source_file}, {locator}]\n"
        excerpt = citation.excerpt.strip()
        remaining = max_characters - used - len(header)
        if remaining <= 0:
            break
        if len(excerpt) > remaining:
            excerpt = excerpt[:remaining].rsplit(" ", 1)[0] + "..."
        entries.append(header + excerpt)
        used += len(header) + len(excerpt)
    return "\n\n".join(entries)
